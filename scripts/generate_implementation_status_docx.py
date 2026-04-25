from __future__ import annotations

import re
import sys
import subprocess
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
BACKEND_DIR = REPO_ROOT / "backend"
FRONTEND_DIR = (
    REPO_ROOT / "frontendv2"
    if (REPO_ROOT / "frontendv2").exists()
    else REPO_ROOT / "frontend"
)
DOCS_DIR = REPO_ROOT / "docs"


@dataclass(frozen=True)
class RouteDef:
    blueprint: str
    path: str
    methods: str
    source_file: str


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_routes(py_file: Path) -> list[RouteDef]:
    """Extract Flask blueprint routes from a Python file by regex."""
    text = _read_text(py_file)

    # Matches: @api_bp.route("/path", methods=["GET"]) or without methods
    # Also supports: @api_bp.route("/a")\n@api_bp.route("/b")
    pat = re.compile(
        r"@(?P<bp>\w+_bp)\.route\(\s*\"(?P<path>[^\"]+)\"(?P<rest>[^)]*)\)"
    )

    def _methods(rest: str) -> str:
        m = re.search(r"methods\s*=\s*\[([^\]]+)\]", rest)
        if not m:
            return "(default)"
        raw = m.group(1)
        vals = re.findall(r"\"(\w+)\"|'(\w+)'", raw)
        methods = [a or b for a, b in vals]
        return ", ".join(methods) if methods else raw.strip()

    routes: list[RouteDef] = []
    for m in pat.finditer(text):
        bp = m.group("bp")
        path = m.group("path")
        rest = m.group("rest") or ""
        routes.append(
            RouteDef(
                blueprint=bp,
                path=path,
                methods=_methods(rest),
                source_file=str(py_file.relative_to(REPO_ROOT)).replace("\\", "/"),
            )
        )

    # Deduplicate (same decorator can appear multiple times)
    uniq: dict[tuple[str, str, str], RouteDef] = {}
    for r in routes:
        uniq[(r.blueprint, r.path, r.methods)] = r
    return sorted(uniq.values(), key=lambda r: (r.blueprint, r.path, r.methods))


def run_check(
    cmd: list[str],
    cwd: Path,
    env: Optional[dict[str, str]] = None,
    timeout_s: int = 300,
) -> tuple[bool, str]:
    """Run a command and return (ok, combined_output)."""
    base_env = dict(**{k: v for k, v in (env or {}).items()})
    try:
        p = subprocess.run(
            cmd,
            cwd=str(cwd),
            env={**dict(**(None or {})), **base_env} if base_env else None,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=timeout_s,
            check=False,
        )
        return p.returncode == 0, (p.stdout or "").strip()
    except Exception as e:
        return False, f"Failed to run {cmd}: {type(e).__name__}: {e}"


def detect_frontend_integration(pages_dir: Path) -> list[list[str]]:
    """Return a per-page status table for React pages."""
    rows: list[list[str]] = []
    index_file = FRONTEND_DIR / "index.html"
    if index_file.exists():
        text = _read_text(index_file)
        calls_api = "/api/" in text
        uses_fetch = "fetch(" in text
        rows.append(
            [
                index_file.name,
                "Integrated" if calls_api else "UI-only (mock/static)",
                "yes" if uses_fetch else "no",
                "no",
                "yes" if calls_api else "no",
            ]
        )
    for js in sorted(pages_dir.glob("*.js")):
        text = _read_text(js)
        calls_api = "http://localhost:5000/api" in text or "/api/" in text
        uses_fetch = "fetch(" in text
        uses_axios = "axios" in text
        status = "Integrated" if calls_api else "UI-only (mock/static)"
        rows.append(
            [
                js.name,
                status,
                "yes" if uses_fetch else "no",
                "yes" if uses_axios else "no",
                "yes" if calls_api else "no",
            ]
        )
    return rows


def build_doc() -> Document:
    doc = Document()

    doc.add_heading("FinGuard Agent — Implementation Status (App Walkthrough)", level=0)
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _p(
        doc,
        "This report is generated by reviewing and executing the application code (backend + frontend), plus running the included test scripts. It maps what is implemented vs what is still missing against the AAS Practice Module Briefing expectations (v4.2).",
    )

    _add_kv_table(
        doc,
        [
            ("Generated", generated_at),
            ("Repo root", str(REPO_ROOT).replace("\\", "/")),
            (
                "Briefing extract",
                "docs/aas_practice_module_briefing_v4.2_full_time.txt",
            ),
            ("Backend", "Flask + SQLAlchemy + CrewAI + ChromaDB"),
            ("Frontend", f"{FRONTEND_DIR.name} + Flask template demo UI"),
        ],
    )

    _h1(doc, "1. What was reviewed (app code, not just docs)")
    _p(
        doc,
        "Backend modules reviewed: app factory/bootstrapping, API routes, auth, audit log, case management, SAR export, agent orchestrator + tools, ML engine, vector store.",
    )
    _p(
        doc,
        f"Frontend modules reviewed: primary UI under {FRONTEND_DIR.name}/ plus the Flask template UI under backend/app/templates/index.html.",
    )

    _h1(doc, "2. Backend capabilities observed in code")
    _h2(doc, "2.1 Portfolio management + analysis")
    _p(
        doc,
        "Implemented: CRUD-style portfolio creation, asset addition/listing, transaction recording, alerts, and AI analysis endpoints.",
    )
    _p(doc, "Evidence: backend/app/routes.py")

    _h2(doc, "2.2 Hybrid ML risk scoring")
    _p(
        doc,
        "Implemented: hybrid scoring (rules + ML models) with rule flags, ML details, hard-block logic, and optional LLM review indicator.",
    )
    _p(doc, "Evidence: backend/ml/risk_scoring_engine.py, backend/app/routes.py")

    _h2(doc, "2.3 Multi-agent orchestration (CrewAI + Groq + RAG)")
    _p(
        doc,
        "Implemented: 9 specialist agents. Orchestrator runs 3 sequential crews to manage rate limits. Agents can use tool functions and RAG context from ChromaDB knowledge base.",
    )
    _p(
        doc,
        "Evidence: backend/agents/crew_orchestrator.py, backend/agents/*, backend/vector_store.py",
    )

    _h2(doc, "2.4 Analyst workflow: cases, audit trail, SAR export")
    _p(
        doc,
        "Implemented: tenant-scoped case queue with state machine transitions; append-only tamper-evident audit log; SAR worksheet export in JSON and PDF.",
    )
    _p(
        doc,
        "Evidence: backend/app/cases.py, backend/app/audit.py, backend/app/sar.py, backend/models/models.py",
    )

    _h2(doc, "2.5 Authentication / authorization")
    _p(
        doc,
        "Implemented: JWT auth with roles (analyst/supervisor/admin). Auth enforcement is configurable (AUTH_ENFORCED). Tenant isolation is applied in cases/audit/SAR.",
    )
    _p(doc, "Evidence: backend/app/auth.py")

    _h1(doc, "3. Frontend capabilities observed in code")

    # React pages table
    pages_dir = FRONTEND_DIR / "src" / "pages"
    if pages_dir.exists() or (FRONTEND_DIR / "index.html").exists():
        _h2(doc, "3.1 React pages: backend integration status")
        _p(
            doc,
            "This table is derived by scanning the active frontend source for API calls (fetch/axios) to /api endpoints.",
        )
        _add_table(
            doc,
            headers=["Page", "Status", "Uses fetch()", "Uses axios", "Calls /api"],
            rows=detect_frontend_integration(pages_dir),
        )
    else:
        _p(doc, f"Frontend source directory not found under {FRONTEND_DIR.name}.")

    _h2(doc, "3.2 Flask template UI")
    _p(
        doc,
        "Implemented: a working demo UI served at / (backend/app/templates/index.html) that calls many backend endpoints (portfolios, transactions risk scoring, AI analysis, search, sentiment, cases).",
    )
    _p(doc, "Evidence: backend/app/__init__.py, backend/app/templates/index.html")

    _h1(doc, "4. Runtime checks (tests executed)")
    _p(
        doc,
        "These checks were executed using the current Python environment. The integration test is run with a placeholder GROQ_API_KEY; it does not make external LLM calls.",
    )

    checks = []

    ok1, out1 = run_check([sys.executable, "-m", "ml.test_engine"], cwd=BACKEND_DIR)
    checks.append(
        ("ML engine test (ml.test_engine)", "PASS" if ok1 else "FAIL", out1[:1500])
    )

    ok2, out2 = run_check([sys.executable, "test_knowledge_base.py"], cwd=BACKEND_DIR)
    checks.append(
        (
            "Knowledge base test (test_knowledge_base.py)",
            "PASS" if ok2 else "FAIL",
            out2[:1500],
        )
    )

    placeholder_key = "gsk_test_key_" + ("x" * 80)
    ok3, out3 = run_check(
        [sys.executable, "test_integration.py"],
        cwd=BACKEND_DIR,
        env={"GROQ_API_KEY": placeholder_key},
    )
    checks.append(
        (
            "Integration test (test_integration.py)",
            "PASS" if ok3 else "FAIL",
            out3[:1500],
        )
    )

    _add_table(
        doc,
        headers=["Check", "Result", "Notes (truncated)"],
        rows=[list(c) for c in checks],
    )

    _h1(doc, "5. Mapping to Practice Module Briefing — Implemented vs Missing")
    _p(
        doc,
        "This section maps the briefing’s ‘Success criteria’ artifact list (system architecture, agent design, XAI/responsible AI, AI security risk register, MLSecOps/LLMSecOps pipeline, tests, UI prototype) to what is present in this repository as working code and runnable checks.",
    )

    mapping_rows = [
        [
            "System architecture document",
            "Implemented (as documentation)",
            "README.md, FINGUARD_SYSTEM_DOCUMENTATION.txt, docs/FinGuard_System_Documentation.docx",
        ],
        [
            "Agent design documentation (reasoning/planning/memory/tools/coordination)",
            "Partially implemented",
            "Agent code exists (backend/agents/*) and orchestration order is implemented (backend/agents/crew_orchestrator.py), but there is no dedicated agent-design report covering patterns/fallbacks/protocols end-to-end.",
        ],
        [
            "Explainable & Responsible AI report (fairness/bias/governance alignment)",
            "Not implemented as a standalone artifact",
            "There is an Explanation agent + XAI knowledge base content, but no report explicitly covering fairness/bias mitigation and governance framework alignment.",
        ],
        [
            "AI security risk register (prompt injection, hallucination, etc.)",
            "Not implemented",
            "No dedicated risk register document found; security mitigations exist in code (auth/audit/tenant scoping) but AI-specific risks are not captured in a register.",
        ],
        [
            "MLSecOps / LLMSecOps pipeline design (CI/CD, automated tests, model versioning, monitoring)",
            "Not implemented",
            "No CI workflow files found (.github/workflows). No pipeline diagram/document in repo.",
        ],
        [
            "Testing artifacts (unit/integration/security)",
            "Partially implemented",
            "Runnable ML+RAG+integration checks exist (backend/ml/test_engine.py, backend/test_knowledge_base.py, backend/test_integration.py). No AI security tests (prompt injection/adversarial) found.",
        ],
        [
            "Simple UI prototype",
            "Implemented",
            "Flask template demo UI is functional; React UI is mostly UI-only except SentimentAnalysis page calling backend.",
        ],
    ]
    _add_table(
        doc, headers=["Briefing item", "Status", "Evidence / notes"], rows=mapping_rows
    )

    _h1(doc, "6. Detailed backlog (what’s left to do)")
    backlog = [
        "Create AI Security Risk Register (table of risks, attack vectors, mitigations, residual risk, owner)",
        "Add AI security tests (prompt injection strings, tool abuse, adversarial inputs) and document expected safe behavior",
        "Write MLSecOps/LLMSecOps pipeline design (CI checks, dependency scanning, model artifact versioning, monitoring/logging plan)",
        "Produce Explainable & Responsible AI report (fairness/bias checks, explainability approach, governance alignment)",
        "Create presentation slides and a demo script aligned to success criteria",
        "(Optional) Increase React frontend-to-backend integration beyond SentimentAnalysis, or position Flask template UI as the primary demo UI",
    ]
    for item in backlog:
        doc.add_paragraph(item, style="List Bullet")

    _h1(doc, "Appendix A: Extracted API endpoints (from code)")
    files_for_routes = [
        BACKEND_DIR / "app" / "routes.py",
        BACKEND_DIR / "app" / "auth.py",
        BACKEND_DIR / "app" / "audit.py",
        BACKEND_DIR / "app" / "cases.py",
        BACKEND_DIR / "app" / "sar.py",
    ]
    all_routes: list[RouteDef] = []
    for f in files_for_routes:
        if f.exists():
            all_routes.extend(extract_routes(f))

    route_rows: list[list[str]] = []
    for r in all_routes:
        prefix = "/api"  # all blueprints in create_app are registered with /api prefix
        full_path = prefix + r.path
        route_rows.append([full_path, r.methods, r.source_file])

    _add_table(doc, headers=["Route", "Methods", "Source"], rows=route_rows[:120])
    if len(route_rows) > 120:
        _p(doc, f"(Truncated: {len(route_rows)} total routes found)")

    return doc


def main() -> int:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCS_DIR / "FinGuard_Implementation_Status_Report.docx"
    doc = build_doc()
    doc.save(out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
