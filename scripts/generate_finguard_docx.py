from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v


def _add_section_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _add_subtitle(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def build_document() -> Document:
    doc = Document()

    doc.add_heading("FinGuard Agent System Document", level=0)
    doc.add_paragraph(
        "This document explains the FinGuard Agent system in plain English. "
        "It covers the agents, the pages, the API endpoints, the order in which agents run, "
        "the user journey, and the technology setup."
    )

    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _add_kv_table(
        doc,
        [
            ("Generated", generated_at),
            ("Scope", "Backend Flask API, CrewAI agents, ML risk engine, ChromaDB, UI pages"),
            ("Primary entry point", "main.py"),
        ],
    )

    _add_section_title(doc, "1. System overview")
    doc.add_paragraph(
        "FinGuard Agent is a financial risk analysis system. It stores portfolios, assets, transactions and alerts in a SQL database. "
        "It runs a hybrid risk scoring engine for transactions and can run a multi agent analysis pipeline for portfolio review. "
        "It also supports an analyst case queue so suspicious activity can be investigated and exported as a SAR worksheet."
    )

    _add_subtitle(doc, "1.1 What the system does")
    doc.add_paragraph("1. Create portfolios and store holdings")
    doc.add_paragraph("2. Record transactions and score risk using rules and ML")
    doc.add_paragraph("3. Auto create alerts and open investigation cases for higher risk activity")
    doc.add_paragraph("4. Run AI analysis across nine specialist agents for portfolio review")
    doc.add_paragraph("5. Provide search over prior AI outputs stored in the vector database")
    doc.add_paragraph("6. Export SAR worksheets as JSON or PDF")

    _add_section_title(doc, "2. Technologies and setup")
    _add_subtitle(doc, "2.1 Main technologies")
    _add_kv_table(
        doc,
        [
            ("Backend API", "Flask and Flask SQLAlchemy"),
            ("Database", "SQLite for relational storage"),
            ("AI orchestration", "CrewAI with sequential tasks"),
            ("LLM provider", "Groq API"),
            ("Vector database", "ChromaDB persistent store"),
            ("ML risk scoring", "Rules plus scikit learn models"),
            ("Frontend", "React app plus a Flask template demo UI"),
        ],
    )

    _add_subtitle(doc, "2.2 Environment configuration")
    doc.add_paragraph(
        "The AI agents require a Groq API key. Set these environment variables in backend/.env or your shell environment."
    )
    _add_kv_table(
        doc,
        [
            ("GROQ_API_KEY", "Required. Used by every agent to call the Groq chat API"),
            ("GROQ_MODEL", "Optional. Defaults to llama-3.3-70b"),
            ("AUTH_ENFORCED", "Optional. If true then a JWT bearer token is required for cases and SAR endpoints"),
            ("JWT_SECRET", "Optional. Secret used to sign JWT tokens"),
            ("CHROMA_PERSIST_DIR", "Optional. Defaults to backend/data/chroma"),
        ],
    )

    _add_subtitle(doc, "2.3 How to start the system")
    doc.add_paragraph("1. From the project root, run: python3 main.py")
    doc.add_paragraph(
        "2. This will install Python dependencies and build the frontend if needed, then start Flask. "
        "The API routes are under /api and the UI is served at the root path."
    )

    _add_section_title(doc, "3. Agents and their functions")
    doc.add_paragraph(
        "There are nine specialist agents. Each agent is a thin wrapper around a Groq LLM call. "
        "Agents also use retrieval from the knowledge base through the vector store when an agent domain is set."
    )

    _add_subtitle(doc, "3.1 Agent list")
    agent_rows = [
        ("Alert Intake Agent", "Processes and categorizes alerts and can enrich transaction alerts with ML risk scores"),
        ("Customer Context Agent", "Builds and summarizes customer profiles and needs"),
        ("Risk Assessment Agent", "Scores portfolio risk and transaction risk with a hybrid rules and ML pipeline and LLM explanation"),
        ("Risk Detection Agent", "Detects fraud patterns and suspicious behavior and assesses market risk exposure"),
        ("Compliance Agent", "Reviews for regulatory and policy compliance and tax reporting issues"),
        ("Portfolio Analysis Agent", "Evaluates allocation and performance and recommends rebalancing actions"),
        ("Market Intelligence Agent", "Produces market sentiment and investment recommendations for symbols"),
        ("Explanation Agent", "Explains risk scores findings and recommendations in clear language"),
        ("Escalation Case Summary Agent", "Decides escalation need and prepares case summaries and escalation packages"),
    ]
    _add_kv_table(doc, agent_rows)

    _add_subtitle(doc, "3.2 Agent interfaces used by the API")
    doc.add_paragraph(
        "The system exposes agent capability through orchestrator methods and API endpoints. "
        "These are the most important entry points used by the application."
    )
    _add_kv_table(
        doc,
        [
            (
                "Full portfolio review",
                "Orchestrator method comprehensive_portfolio_review. Runs nine agents in three sequential crews.",
            ),
            (
                "Quick portfolio recommendation",
                "Orchestrator method quick_portfolio_recommendation. Runs only the Risk Assessment agent in a single sequential CrewAI task.",
            ),
            (
                "Market sentiment",
                "Orchestrator method quick_market_sentiment. Calls Market Intelligence agent analyze_market_sentiment.",
            ),
            (
                "Investment recommendation",
                "Orchestrator method quick_recommendation. Calls Market Intelligence agent generate_investment_recommendation.",
            ),
            (
                "AI explanation for a transaction score",
                "Uses Explanation agent explain_risk_score for human readable insights.",
            ),
        ],
    )

    _add_section_title(doc, "4. Agent execution order")
    doc.add_paragraph(
        "The full portfolio analysis runs in three crews to stay under rate limits. "
        "Each crew runs tasks sequentially. The overall order is fixed."
    )
    doc.add_paragraph("Crew 1 order: Risk Assessment, then Risk Detection, then Compliance")
    doc.add_paragraph("Crew 2 order: Portfolio Analysis, then Market Intelligence, then Customer Context")
    doc.add_paragraph("Crew 3 order: Alert Intake, then Explanation, then Escalation")

    _add_section_title(doc, "5. Pages and their purpose")
    doc.add_paragraph(
        "The repository includes two user interfaces. The React frontend includes several pages, but only the Sentiment page currently calls the backend API. "
        "The Flask template UI is a working demo interface that calls most backend endpoints."
    )

    _add_subtitle(doc, "5.1 React frontend pages")
    _add_kv_table(
        doc,
        [
            ("Dashboard", "Displays sample summary and charts"),
            ("Portfolio", "Shows sample portfolio views"),
            ("Analytics", "Shows sample analytics views"),
            ("Alerts", "Shows sample alert views"),
            ("Sentiment", "Calls the backend to fetch symbols and request sentiment analysis"),
            ("Settings", "Shows sample settings views"),
        ],
    )

    _add_subtitle(doc, "5.2 Flask template demo pages")
    _add_kv_table(
        doc,
        [
            ("Dashboard", "System overview and quick actions"),
            ("Portfolios", "Create portfolio, view portfolio, add assets, view transactions"),
            ("Transaction Risk Analysis", "Record a transaction and score risk and open a case if high risk"),
            ("AI Analysis", "Run full portfolio analysis and request an investment recommendation"),
            ("Alerts", "View alerts"),
            ("Create Alert", "Create an alert"),
            ("Search", "Search prior AI outputs stored in the vector database"),
            ("Sentiment", "Request market sentiment for a symbol"),
            ("Cases", "Analyst case queue and case detail with state transitions and SAR export"),
        ],
    )

    _add_section_title(doc, "6. Page by page interaction logic")
    doc.add_paragraph(
        "This section explains what each page does, which backend endpoints it calls, and which agents are involved. "
        "If a page uses the full portfolio analysis endpoint, the agent order is the full three crew order described earlier."
    )

    _add_subtitle(doc, "6.1 Sentiment page")
    doc.add_paragraph("Purpose: provide AI based market sentiment for one or more symbols")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. GET /api/symbols")
    doc.add_paragraph("2. GET /api/sentiment or GET /api/sentiment/<symbol>")
    doc.add_paragraph("Agents used and order:")
    doc.add_paragraph("1. Market Intelligence agent analyze_market_sentiment")

    _add_subtitle(doc, "6.2 Portfolio management pages")
    doc.add_paragraph("Purpose: create portfolios and manage holdings")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. POST /api/portfolio")
    doc.add_paragraph("2. GET /api/portfolios")
    doc.add_paragraph("3. GET /api/portfolio/<id>")
    doc.add_paragraph("4. POST /api/portfolio/<id>/asset")
    doc.add_paragraph("5. GET /api/portfolio/<id>/assets")
    doc.add_paragraph("Agents used and order: none")

    _add_subtitle(doc, "6.3 Transaction risk analysis page")
    doc.add_paragraph("Purpose: record a transaction and compute a risk score")
    doc.add_paragraph("Backend endpoint:")
    doc.add_paragraph("1. POST /api/portfolio/<id>/transaction")
    doc.add_paragraph("Agents used and order:")
    doc.add_paragraph("1. ML risk engine runs rules and ML models")
    doc.add_paragraph("2. If the risk score is high then an alert is created and a case may be auto opened")
    doc.add_paragraph("3. Optional: the page then calls Explanation agent to generate human readable insights")
    doc.add_paragraph("Backend endpoint for insights:")
    doc.add_paragraph("1. POST /api/transaction/get-ai-insights")
    doc.add_paragraph("Agent used and order for insights:")
    doc.add_paragraph("1. Explanation agent explain_risk_score")

    _add_subtitle(doc, "6.4 Full AI analysis page")
    doc.add_paragraph("Purpose: run the complete nine agent portfolio review")
    doc.add_paragraph("Backend endpoint:")
    doc.add_paragraph("1. POST /api/portfolio/<id>/analyze")
    doc.add_paragraph("Agents used and order:")
    doc.add_paragraph("Crew 1: Risk Assessment, then Risk Detection, then Compliance")
    doc.add_paragraph("Crew 2: Portfolio Analysis, then Market Intelligence, then Customer Context")
    doc.add_paragraph("Crew 3: Alert Intake, then Explanation, then Escalation")

    _add_subtitle(doc, "6.5 Investment recommendation page")
    doc.add_paragraph("Purpose: request an investment recommendation for a selected symbol")
    doc.add_paragraph("Backend endpoint:")
    doc.add_paragraph("1. POST /api/portfolio/<id>/recommendation")
    doc.add_paragraph("Agents used and order:")
    doc.add_paragraph("1. Market Intelligence agent generate_investment_recommendation")

    _add_subtitle(doc, "6.6 Alerts pages")
    doc.add_paragraph("Purpose: view alerts and create alerts")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. GET /api/portfolio/<id>/alerts")
    doc.add_paragraph("2. POST /api/portfolio/<id>/alert")
    doc.add_paragraph("Agents used and order: none")

    _add_subtitle(doc, "6.7 Search page")
    doc.add_paragraph("Purpose: semantic search over prior AI outputs stored in the vector database")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. POST /api/search/analyses")
    doc.add_paragraph("2. POST /api/search/risks")
    doc.add_paragraph("3. POST /api/search/market")
    doc.add_paragraph("Agents used and order: none. This queries ChromaDB.")

    _add_subtitle(doc, "6.8 Cases page")
    doc.add_paragraph("Purpose: analyst queue for investigating suspicious activity")
    doc.add_paragraph("How cases are created:")
    doc.add_paragraph("1. Auto open when a transaction risk score is 55 or higher")
    doc.add_paragraph("2. Or manually opened by an authenticated user")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. GET /api/cases")
    doc.add_paragraph("2. GET /api/cases/<id>")
    doc.add_paragraph("3. POST /api/cases to manually open")
    doc.add_paragraph("4. POST /api/cases/<id>/assign")
    doc.add_paragraph("5. POST /api/cases/<id>/notes")
    doc.add_paragraph("6. POST /api/cases/<id>/transition")
    doc.add_paragraph("7. POST /api/cases/<id>/analyze")
    doc.add_paragraph("8. GET /api/cases/<id>/customer-360")
    doc.add_paragraph("Agents used and order for re run analysis from a case:")
    doc.add_paragraph("1. risk action uses Risk Detection agent")
    doc.add_paragraph("2. compliance action uses Compliance agent")
    doc.add_paragraph("3. portfolio action uses Portfolio Analysis agent")
    doc.add_paragraph("4. recommendation action uses a single Risk Assessment agent crew")

    _add_subtitle(doc, "6.9 SAR export")
    doc.add_paragraph("Purpose: export a SAR worksheet for a case")
    doc.add_paragraph("Backend endpoints:")
    doc.add_paragraph("1. GET /api/sar/<case_id>.json")
    doc.add_paragraph("2. GET /api/sar/<case_id>.pdf")
    doc.add_paragraph("Agents used and order: none")

    _add_section_title(doc, "7. User journey")
    doc.add_paragraph(
        "A typical end to end workflow for an analyst or portfolio user looks like this."
    )
    doc.add_paragraph("1. Create a portfolio")
    doc.add_paragraph("2. Add assets to the portfolio")
    doc.add_paragraph("3. Record transactions")
    doc.add_paragraph("4. Review the computed risk score")
    doc.add_paragraph("5. If the score is high, review the auto created case in the case queue")
    doc.add_paragraph("6. Re run focused analysis inside the case when needed")
    doc.add_paragraph("7. Run full portfolio AI analysis for a broad review")
    doc.add_paragraph("8. Search prior analyses and risk outputs")
    doc.add_paragraph("9. Export a SAR worksheet when required")

    _add_section_title(doc, "8. API summary")
    doc.add_paragraph("This is a quick reference list of the core API endpoints.")
    api_rows = [
        ("Symbols", "GET /api/symbols and GET /api/symbols/sectors"),
        ("Portfolio", "POST /api/portfolio, GET /api/portfolios, GET /api/portfolio/<id>"),
        ("Assets", "POST /api/portfolio/<id>/asset, GET /api/portfolio/<id>/assets"),
        ("Transactions", "POST /api/portfolio/<id>/transaction, GET /api/portfolio/<id>/transactions"),
        ("Standalone risk scoring", "POST /api/transaction/score-risk"),
        ("AI insights", "POST /api/transaction/get-ai-insights"),
        ("Full portfolio analysis", "POST /api/portfolio/<id>/analyze"),
        ("Investment recommendation", "POST /api/portfolio/<id>/recommendation"),
        ("Sentiment", "GET /api/sentiment or GET /api/sentiment/<symbol>"),
        ("Alerts", "GET /api/portfolio/<id>/alerts and POST /api/portfolio/<id>/alert"),
        ("Search", "POST /api/search/analyses, /api/search/risks, /api/search/market"),
        ("Cases", "GET and POST /api/cases plus case detail and mutation endpoints"),
        ("SAR", "GET /api/sar/<case_id>.json and /api/sar/<case_id>.pdf"),
    ]
    _add_kv_table(doc, api_rows)

    return doc


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "docs" / "FinGuard_System_Documentation.docx"

    doc = build_document()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
